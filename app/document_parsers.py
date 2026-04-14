from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

import docx
from docling.document_converter import DocumentConverter
from pypdf import PdfReader
from unstructured.partition.auto import partition
from unstructured.partition.pdf import partition_pdf


@dataclass
class ParsedPage:
    page_number: int
    text: str


@dataclass
class FieldValue:
    name: str
    value: str
    source_text: str
    page_number: int
    aliases: list[str] = field(default_factory=list)


@dataclass
class ParsedBlock:
    block_id: str
    page_number: int
    section_name: str
    chunk_type: str
    text: str
    aliases: list[str] = field(default_factory=list)


@dataclass
class ParsedDocument:
    filename: str
    doc_type: str
    pages: list[ParsedPage]
    markdown: str
    blocks: list[ParsedBlock]
    fields: dict[str, FieldValue]


_DOCLING_CONVERTER: DocumentConverter | None = None


def parse_document(path: Path) -> ParsedDocument:
    pages = _parse_pages(path)
    markdown = _parse_markdown(path)
    doc_type = detect_doc_type(path.name, markdown or "\n".join(page.text for page in pages))
    blocks = build_blocks(markdown, pages)
    fields = extract_fields(doc_type, markdown, pages, blocks)
    return ParsedDocument(
        filename=path.name,
        doc_type=doc_type,
        pages=pages,
        markdown=markdown,
        blocks=blocks,
        fields=fields,
    )


def detect_doc_type(filename: str, text: str) -> str:
    head = f"{filename}\n{text[:2000]}".lower()
    name = filename.lower()
    if "bill of lading" in head:
        return "bill_of_lading"
    if "carrier-rc" in name or "carrier rate" in head:
        return "carrier_rate_confirmation"
    if "shipper-rc" in name or "customer rate" in head or "shipper" in name:
        return "shipper_rate_confirmation"
    if "invoice" in head:
        return "invoice"
    if "instruction" in head:
        return "shipment_instructions"
    return "unknown"


def _parse_pages(path: Path) -> list[ParsedPage]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        return [
            ParsedPage(page_number=index + 1, text=_clean_text(page.extract_text() or ""))
            for index, page in enumerate(reader.pages)
        ]
    if suffix == ".docx":
        document = docx.Document(str(path))
        text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
        return [ParsedPage(page_number=1, text=_clean_text(text))]
    if suffix == ".txt":
        return [ParsedPage(page_number=1, text=_clean_text(path.read_text(encoding="utf-8")))]
    raise ValueError("Unsupported file type")


def _parse_markdown(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        try:
            return _parse_with_docling(path)
        except Exception:
            return _parse_with_unstructured_pdf(path)
    if suffix in {".docx", ".txt"}:
        try:
            elements = partition(filename=str(path))
            lines = [str(element).strip() for element in elements if str(element).strip()]
            return _clean_text("\n\n".join(lines))
        except Exception:
            return "\n\n".join(page.text for page in _parse_pages(path))
    raise ValueError("Unsupported file type")


def _parse_with_docling(path: Path) -> str:
    global _DOCLING_CONVERTER
    if _DOCLING_CONVERTER is None:
        _DOCLING_CONVERTER = DocumentConverter()
    result = _DOCLING_CONVERTER.convert(path)
    markdown = result.document.export_to_markdown()
    markdown = _clean_text(markdown.replace("<!-- image -->", ""))
    if len(markdown) < 200:
        raise ValueError("Docling output too small")
    return markdown


def _parse_with_unstructured_pdf(path: Path) -> str:
    elements = partition_pdf(
        filename=str(path),
        strategy="fast",
        infer_table_structure=True,
    )
    lines = [str(element).strip() for element in elements if str(element).strip()]
    return _clean_text("\n\n".join(lines))


def build_blocks(markdown: str, pages: list[ParsedPage]) -> list[ParsedBlock]:
    blocks: list[ParsedBlock] = []
    lines = markdown.splitlines()
    current_heading = "document"
    table_counter = 0
    paragraph_counter = 0
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.startswith("#"):
            current_heading = line.lstrip("#").strip().lower().replace(" ", "_")
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 2:
                table_counter += 1
                raw = "\n".join(table_lines)
                blocks.extend(_table_to_blocks(raw, current_heading, table_counter, pages))
            continue
        paragraph_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("|") and not lines[i].startswith("#"):
            paragraph_lines.append(lines[i].rstrip())
            i += 1
        paragraph_text = _clean_text("\n".join(paragraph_lines))
        if paragraph_text:
            paragraph_counter += 1
            blocks.append(
                ParsedBlock(
                    block_id=f"text-{paragraph_counter}",
                    page_number=_page_for_text(paragraph_text, pages),
                    section_name=current_heading,
                    chunk_type="text",
                    text=paragraph_text,
                    aliases=_aliases_for_text(paragraph_text, current_heading),
                )
            )
    return blocks


def extract_fields(
    doc_type: str,
    markdown: str,
    pages: list[ParsedPage],
    blocks: list[ParsedBlock],
) -> dict[str, FieldValue]:
    fields: dict[str, FieldValue] = {}
    full_text = "\n\n".join(page.text for page in pages)
    combined = f"{markdown}\n\n{full_text}"

    _add_regex_field(fields, "shipment_id", r"(?:Load ID|Reference ID)\s+([A-Z0-9-]+)", combined, pages, ["reference id", "load id", "shipment id"])
    _add_regex_field(fields, "weight", r"([0-9][0-9,.\s]*lbs)", combined, pages, ["weight"])
    _add_regex_field(fields, "mode", r"Load Type\s+([A-Z]+)", combined, pages, ["mode", "load type"])
    _add_regex_field(fields, "currency", r"\b(USD)\b", combined, pages, ["currency", "usd"])
    _add_regex_field(fields, "pickup_datetime", r"(?:Ship Date|Shipping Date)\s*\|?\s*([0-9-]+(?:\s+[0-9:]+)?)", combined, pages, ["pickup", "shipping date", "ship date"])
    _add_regex_field(fields, "delivery_datetime", r"Delivery Date\s*\|?\s*([0-9-]+(?:\s+[0-9:]+)?)", combined, pages, ["delivery", "delivery date"])

    table_blocks = [block for block in blocks if block.chunk_type == "table_row"]
    for block in table_blocks:
        lower = block.section_name
        text = block.text
        page = block.page_number
        if "shipper" in text.lower() and "consignee" in text.lower():
            continue
        if lower == "carrier_details":
            _add_from_table_text(fields, "carrier_name", text, r"carrier: ([^|]+)", page, ["carrier", "carrier name"])
            _add_from_table_text(fields, "equipment_type", text, r"equipment: ([^|]+)", page, ["equipment", "equipment type"])
            _add_from_table_text(fields, "rate", text, r"agreed amount \(usd\): \$?([0-9.]+)", page, ["rate", "carrier rate", "agreed amount"])
        if lower == "customer_details":
            _add_from_table_text(fields, "rate", text, r"agreed amount \(usd\): \$?([0-9.]+)", page, ["rate", "agreed amount"])

    if "shipper" not in fields or "consignee" not in fields:
        shipper, consignee, source_text = _extract_bol_parties(markdown)
        if shipper:
            fields["shipper"] = FieldValue("shipper", shipper, source_text, _page_for_text(source_text, pages), ["shipper", "consignor"])
        if consignee:
            fields["consignee"] = FieldValue("consignee", consignee, source_text, _page_for_text(source_text, pages), ["consignee", "receiver"])

    if "shipper" not in fields:
        pickup_name, pickup_source = _extract_stop_name(markdown, "Pickup")
        if pickup_name:
            fields["shipper"] = FieldValue("shipper", pickup_name, pickup_source, _page_for_text(pickup_source, pages), ["shipper", "pickup"])
    if "consignee" not in fields:
        drop_name, drop_source = _extract_stop_name(markdown, "Drop")
        if drop_name:
            fields["consignee"] = FieldValue("consignee", drop_name, drop_source, _page_for_text(drop_source, pages), ["consignee", "drop"])

    pickup_dt = _extract_stop_datetime(markdown, "Pickup")
    if pickup_dt:
        fields["pickup_datetime"] = FieldValue(
            "pickup_datetime",
            pickup_dt,
            pickup_dt,
            _page_for_text(pickup_dt, pages),
            ["pickup", "pickup date", "shipping date", "appointment"],
        )
    delivery_dt = _extract_stop_datetime(markdown, "Drop")
    if delivery_dt:
        fields["delivery_datetime"] = FieldValue(
            "delivery_datetime",
            delivery_dt,
            delivery_dt,
            _page_for_text(delivery_dt, pages),
            ["delivery", "drop", "delivery date", "appointment"],
        )

    if doc_type == "bill_of_lading" and "rate" not in fields:
        cod_value = re.search(r"COD Value\s+\$?([0-9.]+)\s*USD", full_text, flags=re.IGNORECASE)
        if cod_value:
            fields["rate"] = FieldValue("rate", cod_value.group(1), cod_value.group(0), _page_for_text(cod_value.group(0), pages), ["rate", "cod value"])

    _add_labeled_field(
        fields,
        "shipment_id",
        [r"^\s*Shipment(?:\s+ID)?\s*[:\-]\s*([A-Z0-9-]+)\s*$"],
        combined,
        pages,
        ["shipment id"],
    )
    _add_labeled_field(
        fields,
        "shipper",
        [r"^\s*Shipper\s*[:\-]\s*(.+?)\s*$"],
        combined,
        pages,
        ["shipper", "consignor"],
    )
    _add_labeled_field(
        fields,
        "consignee",
        [r"^\s*Consignee\s*[:\-]\s*(.+?)\s*$"],
        combined,
        pages,
        ["consignee", "receiver"],
    )
    _add_labeled_field(
        fields,
        "pickup_datetime",
        [
            r"^\s*Pickup(?:\s+Datetime|\s+Date\s+Time|\s+Date)?\s*[:\-]\s*(.+?)\s*$",
            r"^\s*Ship(?:ping)?\s+Date\s*[:\-]\s*(.+?)\s*$",
        ],
        combined,
        pages,
        ["pickup", "pickup date", "shipping date"],
    )
    _add_labeled_field(
        fields,
        "delivery_datetime",
        [
            r"^\s*Delivery(?:\s+Datetime|\s+Date\s+Time|\s+Date)?\s*[:\-]\s*(.+?)\s*$",
            r"^\s*Drop(?:\s+Datetime|\s+Date)?\s*[:\-]\s*(.+?)\s*$",
        ],
        combined,
        pages,
        ["delivery", "drop", "delivery date"],
    )
    _add_labeled_field(
        fields,
        "equipment_type",
        [r"^\s*Equipment(?:\s+Type)?\s*[:\-]\s*(.+?)\s*$"],
        combined,
        pages,
        ["equipment", "equipment type"],
    )
    _add_labeled_field(
        fields,
        "mode",
        [r"^\s*(?:Mode|Load\s+Type)\s*[:\-]\s*(.+?)\s*$"],
        combined,
        pages,
        ["mode", "load type"],
    )
    _add_labeled_field(
        fields,
        "rate",
        [r"^\s*(?:Carrier\s+Rate|Agreed\s+Amount|Rate)\s*[:\-]\s*\$?([0-9.,]+)\s*(?:[A-Z]{3})?\s*$"],
        combined,
        pages,
        ["rate", "carrier rate", "agreed amount"],
    )
    _add_labeled_field(
        fields,
        "currency",
        [r"^\s*Currency\s*[:\-]\s*([A-Z]{3})\s*$"],
        combined,
        pages,
        ["currency"],
        transform=lambda value: value.upper(),
    )
    _add_labeled_field(
        fields,
        "weight",
        [r"^\s*Weight\s*[:\-]\s*([0-9][0-9,.\s]*lbs)\s*$"],
        combined,
        pages,
        ["weight"],
    )
    _add_labeled_field(
        fields,
        "carrier_name",
        [r"^\s*Carrier(?:\s+Name)?\s*[:\-]\s*(.+?)\s*$"],
        combined,
        pages,
        ["carrier", "carrier name"],
    )

    return fields


def _table_to_blocks(raw_table: str, heading: str, table_counter: int, pages: list[ParsedPage]) -> list[ParsedBlock]:
    rows = [line for line in raw_table.splitlines() if line.strip()]
    if len(rows) < 2:
        return []
    headers = [_strip_cell(cell) for cell in rows[0].split("|")[1:-1]]
    data_rows = rows[2:] if len(rows) > 2 else []
    blocks: list[ParsedBlock] = []
    if not data_rows:
        return blocks
    for row_index, row in enumerate(data_rows, start=1):
        values = [_strip_cell(cell) for cell in row.split("|")[1:-1]]
        if len(values) != len(headers):
            continue
        pairs = [f"{header}: {value}" for header, value in zip(headers, values) if value]
        row_text = " | ".join(pairs)
        page_number = _page_for_text(" ".join(values), pages)
        blocks.append(
            ParsedBlock(
                block_id=f"table-{table_counter}-row-{row_index}",
                page_number=page_number,
                section_name=heading,
                chunk_type="table_row",
                text=row_text,
                aliases=[header.lower() for header in headers if header],
            )
        )
    return blocks


def _extract_bol_parties(markdown: str) -> tuple[str | None, str | None, str]:
    pattern = re.compile(
        r"\|\s*Shipper\s*\|\s*Consignee\s*\|.*?\n\|[-| ]+\|\n\|\s*(.*?)\s*\|\s*(.*?)\s*\|",
        flags=re.IGNORECASE | re.DOTALL,
    )
    match = pattern.search(markdown)
    if not match:
        return None, None, ""
    shipper_cell, consignee_cell = match.groups()
    shipper = _cleanup_party_name(shipper_cell)
    consignee = _cleanup_party_name(consignee_cell)
    return shipper, consignee, match.group(0)


def _extract_stop_name(markdown: str, label: str) -> tuple[str | None, str]:
    pattern = re.compile(
        rf"\|\s*\d+\s*\|\s*{label}\s*\|.*?\n\|\s*\|\s*(.*?)\s*\|",
        flags=re.IGNORECASE | re.DOTALL,
    )
    match = pattern.search(markdown)
    if not match:
        return None, ""
    source = match.group(0)
    return _cleanup_party_name(match.group(1)), source


def _extract_stop_datetime(markdown: str, label: str) -> str | None:
    if label.lower() == "pickup":
        pattern = re.compile(r"Shipping Date\s+([0-9-]+).*?Appointment\s+([0-9:-]+)", flags=re.IGNORECASE | re.DOTALL)
    else:
        pattern = re.compile(r"Delivery Date\s+([0-9-]+).*?Appointment\s+([0-9:-]+|-)", flags=re.IGNORECASE | re.DOTALL)
    match = pattern.search(markdown)
    if not match:
        return None
    date_part, appointment = match.groups()
    appointment = appointment.strip()
    if appointment and appointment != "-":
        return f"{date_part} {appointment}"
    return date_part


def _add_regex_field(
    fields: dict[str, FieldValue],
    name: str,
    pattern: str,
    text: str,
    pages: list[ParsedPage],
    aliases: list[str],
) -> None:
    if name in fields:
        return
    match = re.search(pattern, text, flags=re.IGNORECASE)
    if not match:
        return
    value = match.group(1).strip()
    fields[name] = FieldValue(name, value, match.group(0), _page_for_text(match.group(0), pages), aliases)


def _add_from_table_text(
    fields: dict[str, FieldValue],
    name: str,
    text: str,
    pattern: str,
    page_number: int,
    aliases: list[str],
) -> None:
    if name in fields:
        return
    match = re.search(pattern, text, flags=re.IGNORECASE)
    if not match:
        return
    fields[name] = FieldValue(name, match.group(1).strip(), text, page_number, aliases)


def _add_labeled_field(
    fields: dict[str, FieldValue],
    name: str,
    patterns: list[str],
    text: str,
    pages: list[ParsedPage],
    aliases: list[str],
    transform=None,
) -> None:
    if name in fields:
        return
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE | re.MULTILINE)
        if not match:
            continue
        value = match.group(1).strip()
        value = transform(value) if transform else _clean_labeled_value(value)
        if not value:
            continue
        source = match.group(0).strip()
        fields[name] = FieldValue(name, value, source, _page_for_text(source, pages), aliases)
        return


def _page_for_text(text: str, pages: list[ParsedPage]) -> int:
    snippet = _clean_text(text)[:120]
    if not snippet:
        return 1
    for page in pages:
        if snippet[:40] in page.text or _cleanup_party_name(snippet) in page.text:
            return page.page_number
    return 1


def _cleanup_party_name(value: str) -> str:
    value = re.sub(r"^\d+\.\s*", "", value.strip())
    value = value.split(",")[0].strip()
    value = re.split(r"\s+\d", value, maxsplit=1)[0].strip()
    value = re.split(r"\s+Los Angeles\b", value, maxsplit=1)[0].strip()
    value = re.split(r"\s+International Airport\b", value, maxsplit=1)[0].strip()
    value = re.split(r"\s+Cherry Avenue\b", value, maxsplit=1)[0].strip()
    value = re.split(r"\s+Fontana\b", value, maxsplit=1)[0].strip()
    tokens = value.split()
    if tokens and len(tokens[0]) <= 12:
        return tokens[0].strip()
    return value or value


def _clean_labeled_value(value: str) -> str:
    value = _clean_text(value)
    value = re.sub(r"\s+[|]\s*$", "", value)
    value = re.sub(r"\s{2,}", " ", value)
    return value.strip(" -:\t")


def _aliases_for_text(text: str, heading: str) -> list[str]:
    aliases = [heading]
    lower = text.lower()
    for term in [
        "shipper", "consignee", "carrier", "rate", "pickup",
        "delivery", "drop", "appointment", "weight", "equipment",
        "reference id", "load id", "customer", "commodity",
    ]:
        if term in lower or term in heading:
            aliases.append(term)
    return sorted(set(alias for alias in aliases if alias))


def _strip_cell(value: str) -> str:
    return _clean_text(value.replace("\\|", "|"))


def _clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = text.replace("Conﬁrmation", "Confirmation")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
